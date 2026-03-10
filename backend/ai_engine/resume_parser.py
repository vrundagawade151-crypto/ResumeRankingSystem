"""Resume parsing using NLP - extracts text from PDF/DOCX and parses skills, education, experience."""
import re
import os
from pathlib import Path

# PDF extraction
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX extraction
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# NLTK for NLP
import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk import pos_tag


class ResumeParser:
    """Parse resume files and extract structured data using NLP."""

    # Common skill keywords for extraction
    SKILL_KEYWORDS = [
        'python', 'java', 'javascript', 'react', 'node', 'sql', 'mysql', 'mongodb',
        'aws', 'docker', 'kubernetes', 'git', 'html', 'css', 'typescript', 'angular',
        'vue', 'django', 'flask', 'fastapi', 'machine learning', 'deep learning',
        'nlp', 'data analysis', 'excel', 'power bi', 'tableau', 'agile', 'scrum',
        'rest api', 'graphql', 'redis', 'postgresql', 'linux', 'ci/cd', 'jenkins',
        'figma', 'ui/ux', 'photoshop', 'illustrator', 'project management',
        'communication', 'leadership', 'problem solving', 'teamwork', 'analytical',
        'c++', 'c#', '.net', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
        'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'spacy', 'nltk'
    ]

    # Education degree patterns
    EDUCATION_PATTERNS = [
        r'\b(bachelor|b\.?s\.?|b\.?tech|b\.?e\.?|b\.?a\.?)\b',
        r'\b(master|m\.?s\.?|m\.?tech|m\.?e\.?|m\.?a\.?|mba|m\.?b\.?a\.?)\b',
        r'\b(ph\.?d|doctorate|doctoral)\b',
        r'\b(high school|secondary school|diploma|associate)\b',
        r'\b(b\.?com|b\.?sc|m\.?sc|b\.?ca|m\.?ca)\b',
        r'\b(university|college|institute|institution)\b',
        r'\b(computer science|information technology|engineering|electrical|mechanical|civil)\b',
        r'\b(\d{4})\s*[-–]\s*(\d{4}|\bpresent\b|\bcurrent\b)\b'
    ]

    # Experience patterns
    EXPERIENCE_PATTERNS = [
        r'\b(\d+)\+?\s*(years?|yrs?)\s*(of\s*)?(experience|exp\.?)\b',
        r'\b(experience|exp\.?)\s*[:\-]?\s*(\d+)\+?\s*(years?|yrs?)\b',
        r'\b(intern|internship|full[- ]?time|part[- ]?time|contract|freelance)\b',
        r'\b(senior|junior|lead|principal|associate|entry[- ]?level)\b',
        r'\b(software engineer|developer|engineer|analyst|manager|director)\b',
        r'\b(at|@)\s+([A-Za-z0-9\s&\.]+?)(?:\s*[-–]\s*|\s*$|\n)',
        r'\b(company|corporation|inc\.?|ltd\.?|llc)\b'
    ]

    def __init__(self):
        self.stop_words = set(stopwords.words('english'))

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return text.strip()

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")

    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")

    def extract_name(self, text: str) -> str:
        """Extract candidate name - typically at the start of resume."""
        lines = text.split('\n')
        for line in lines[:10]:  # Name usually in first 10 lines
            line = line.strip()
            if len(line) > 2 and len(line) < 50:
                # Skip lines that look like emails or phone numbers
                if '@' in line or re.search(r'\d{10}', line):
                    continue
                # Check if it looks like a name (capitalized words, 2-4 parts)
                words = line.split()
                if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if len(w) > 1):
                    return line
        return "Unknown"

    def extract_skills(self, text: str) -> str:
        """Extract skills from resume text using keyword matching and NLP."""
        text_lower = text.lower()
        found_skills = set()
        
        # Direct keyword matching
        for skill in self.SKILL_KEYWORDS:
            if skill in text_lower:
                found_skills.add(skill.title())
        
        # Extract capitalized phrases that might be skills (between sections)
        lines = text.split('\n')
        for line in lines:
            # Look for comma-separated or bullet-separated skill lists
            parts = re.split(r'[,•\-\|;:]', line)
            for part in parts:
                part = part.strip()
                if 2 <= len(part) <= 30 and part.lower() in [s.lower() for s in self.SKILL_KEYWORDS]:
                    found_skills.add(part.title())
        
        # Use NLTK to find noun phrases that might be skills
        words = word_tokenize(text_lower)
        tagged = pos_tag(words)
        for i, (word, pos) in enumerate(tagged):
            if pos in ['NN', 'NNS', 'NNP', 'NNPS'] and len(word) > 2:
                if word.lower() in [s.lower() for s in self.SKILL_KEYWORDS]:
                    found_skills.add(word.title())
        
        return ', '.join(sorted(found_skills)) if found_skills else "Not specified"

    def extract_education(self, text: str) -> str:
        """Extract education information from resume."""
        text_lower = text.lower()
        education_parts = []
        
        for pattern in self.EDUCATION_PATTERNS:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                # Get surrounding context
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 100)
                context = text[start:end]
                # Clean and add
                clean = ' '.join(context.split())[:80]
                if clean and clean not in str(education_parts):
                    education_parts.append(clean)
        
        # Fallback: find education section
        edu_section = re.search(
            r'(?:education|academic|qualification)s?\s*[:\-]*(.*?)(?=\n\n|\b(?:experience|work|skills)\b|$)',
            text_lower, re.DOTALL | re.IGNORECASE
        )
        if edu_section:
            edu_text = edu_section.group(1).strip()[:300]
            education_parts.append(edu_text)
        
        return ' | '.join(education_parts[:3]) if education_parts else "Not specified"

    def extract_experience(self, text: str) -> str:
        """Extract experience information from resume."""
        text_lower = text.lower()
        experience_parts = []
        
        # Years of experience
        years_match = re.search(r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp\.?)', text_lower)
        if years_match:
            experience_parts.append(f"{years_match.group(1)}+ years experience")
        
        # Job titles
        for pattern in self.EXPERIENCE_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                context = text[max(0, match.start()-20):min(len(text), match.end()+60)]
                clean = ' '.join(context.split())[:100]
                if clean and len(clean) > 10:
                    experience_parts.append(clean)
        
        # Experience section
        exp_section = re.search(
            r'(?:experience|work\s*history|employment)s?\s*[:\-]*(.*?)(?=\n\n|\b(?:education|skills)\b|$)',
            text_lower, re.DOTALL | re.IGNORECASE
        )
        if exp_section:
            exp_text = exp_section.group(1).strip()[:400]
            experience_parts.append(exp_text)
        
        return ' | '.join(experience_parts[:3]) if experience_parts else "Not specified"

    def parse(self, file_path: str) -> dict:
        """Parse resume file and return extracted data."""
        raw_text = self.extract_text(file_path)
        
        return {
            'raw_text': raw_text[:5000],  # Limit stored text
            'name': self.extract_name(raw_text),
            'skills': self.extract_skills(raw_text),
            'education': self.extract_education(raw_text),
            'experience': self.extract_experience(raw_text)
        }
